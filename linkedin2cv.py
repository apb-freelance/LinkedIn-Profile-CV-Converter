import tarfile
import os
import sqlite3
import pandas as pd
from fpdf import FPDF
from docx import Document

def extract_tar(file_path, extract_to):
	"""Extrae un archivo TAR en un directorio específico con manejo de filtrado en Python 3.14."""
	with tarfile.open(file_path, "r") as tar:
		tar.extractall(extract_to, filter="data")  # Añadido filtro para compatibilidad con Python 3.14
	return os.path.join(extract_to, "profile")

def load_csv_to_sqlite(db_name, folder_path):
	"""Carga archivos CSV en una base de datos SQLite."""
	conn = sqlite3.connect(db_name)
	csv_files = [f for f in os.listdir(folder_path) if f.endswith(".csv")]
	for file in csv_files:
		table_name = os.path.splitext(file)[0]
		file_path = os.path.join(folder_path, file)
		df = pd.read_csv(file_path)
		df.to_sql(table_name, conn, index=False, if_exists="replace")
	return conn

def data_format(data = ""):
	data = data.replace(" · ", "\n\t· ")
	return data

def get_tables():
	tables = tables = ["PhoneNumbers", 
					"Email Addresses", 
					"Profile", 
					"Positions", 
					"Education", 
					"Projects", 
					"Languages", 
					"Skills"]
	return tables

def generate_cv_pdf(conn, output_path):
	"""Genera un CV en PDF basado en la base de datos SQLite."""
	pdf = FPDF()
	pdf.set_auto_page_break(auto=True, margin=15)
	pdf.add_page()
	pdf.set_font("Arial", style="B", size=16)

	tables = get_tables()
	for table in tables:
		data = pd.read_sql_query(f"SELECT * FROM '{table}';", conn)
		if not data.empty:
			data = data_format(data)
			pdf.set_font("Arial", style="B", size=14)
			pdf.cell(200, 10, table.replace('_', ' '), ln=True)
			pdf.set_font("Arial", size=12)
			for _, row in data.iterrows():
				pdf.multi_cell(0, 10, ", ".join(map(str, row.dropna().tolist())))
				pdf.ln(5)
	
	pdf.output(output_path)

def generate_cv_txt(conn, output_path):
	"""Genera un CV en formato TXT."""
	with open(output_path, "w") as file:
		tables = get_tables()
		for table in tables:
			data = pd.read_sql_query(f"SELECT * FROM '{table}';", conn)
			if not data.empty:
				data = data_format(data)
				file.write(f"{table.replace('_', ' ')}:\n\n")
				for _, row in data.iterrows():
					file.write("\t" + data_format(", ".join(map(str, row.dropna().tolist())) + "\n"))
				file.write("\n")

def generate_cv_docx(conn, output_path):
	"""Genera un CV en formato DOCX directamente desde la base de datos."""
	doc = Document()
	doc.add_heading("Curriculum Vitae", level=1)
	
	tables = get_tables()
	for table in tables:
		data = pd.read_sql_query(f"SELECT * FROM '{table}';", conn)
		if not data.empty:
			data = data_format(data)
			doc.add_heading(table.replace('_', ' '), level=2)
			for _, row in data.iterrows():
				doc.add_paragraph(", ".join(map(str, row.dropna().tolist())))
	
	doc.save(output_path)

def package_cv(files, tar_output):
	"""Empaqueta los archivos de CV en un archivo TAR."""
	with tarfile.open(tar_output, "w") as tar:
		for file in files:
			tar.add(file, arcname=os.path.basename(file))

if __name__ == "__main__":
	tar_file = "profile.tar"
	extract_folder = "./extracted"
	db_name = ":memory:"
	cv_pdf = "cv_generated.pdf"
	cv_txt = "cv_generated.txt"
	cv_docx = "cv_generated.docx"
	tar_output = "cv_generated.tar"
	
	extracted_path = extract_tar(tar_file, extract_folder)
	conn = load_csv_to_sqlite(db_name, extracted_path)
	
	generate_cv_pdf(conn, cv_pdf)
	generate_cv_txt(conn, cv_txt)
	generate_cv_docx(conn, cv_docx)
	
